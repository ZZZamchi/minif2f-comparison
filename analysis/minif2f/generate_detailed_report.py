#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成详细的 miniF2F v1 vs v2 比较报告
包含具体示例，并生成 Word 文档
"""

import json
import sys
from pathlib import Path
from collections import defaultdict
from typing import Dict, List, Any, Optional

# 设置输出编码（Windows 兼容）
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARN] python-docx 未安装，将生成 Markdown 格式。安装: pip install python-docx")


class DetailedReportGenerator:
    """详细报告生成器"""
    
    def __init__(self, v1_path: str, v2c_path: str, v2s_path: str):
        self.v1_path = v1_path
        self.v2c_path = v2c_path
        self.v2s_path = v2s_path
        
        self.v1_data = {}
        self.v2c_data = {}
        self.v2s_data = {}
        
    def load_datasets(self):
        """加载所有数据集"""
        print("正在加载数据集...")
        
        with open(self.v1_path, 'r', encoding='utf-8') as f:
            v1_list = json.load(f)
            self.v1_data = {item['name']: item for item in v1_list}
        print(f"[OK] v1: {len(self.v1_data)} 条记录")
        
        with open(self.v2c_path, 'r', encoding='utf-8') as f:
            v2c_list = json.load(f)
            self.v2c_data = {item['name']: item for item in v2c_list}
        print(f"[OK] v2c: {len(self.v2c_data)} 条记录")
        
        with open(self.v2s_path, 'r', encoding='utf-8') as f:
            v2s_list = json.load(f)
            self.v2s_data = {item['name']: item for item in v2s_list}
        print(f"[OK] v2s: {len(self.v2s_data)} 条记录")
    
    def find_comparison_examples(self) -> List[Dict]:
        """找到适合比较的示例"""
        examples = []
        
        # 找出共同的问题
        common_names = set(self.v1_data.keys()) & set(self.v2c_data.keys()) & set(self.v2s_data.keys())
        
        # 优先选择 AMC 问题（最能体现差异）
        amc_names = [name for name in common_names if 'amc' in name.lower()]
        
        # 选择几个代表性的问题
        selected_names = []
        
        # 1. 选择一个 AMC 选择题（v1 和 v2s 直接给出答案，v2c 包含所有选项）
        for name in amc_names[:20]:
            v1_item = self.v1_data[name]
            v2c_item = self.v2c_data[name]
            # 确保 v2c 包含选项，v1 直接给出答案
            if ('Show that it is' in v1_item.get('informal_statement', '') and
                'Prove that it is equal to one of the following options' in v2c_item.get('informal_statement', '')):
                selected_names.append(name)
                break
        
        # 2. 选择另一个 AMC 问题（确保有差异）
        for name in amc_names[1:20]:
            if name not in selected_names:
                v1_item = self.v1_data[name]
                v2c_item = self.v2c_data[name]
                if (v1_item.get('informal_statement', '') != v2c_item.get('informal_statement', '') or
                    v1_item.get('formal_statement', '') != v2c_item.get('formal_statement', '')):
                    selected_names.append(name)
                    break
        
        # 3. 选择一个非选择题（如 mathd, aime, imo 等）
        non_amc = [name for name in common_names if 'amc' not in name.lower()]
        for name in non_amc[:10]:
            if name not in selected_names:
                selected_names.append(name)
                break
        
        # 4. 再选一个 AMC 问题
        for name in amc_names[2:20]:
            if name not in selected_names and len(selected_names) < 5:
                selected_names.append(name)
                break
        
        # 生成示例
        for name in selected_names[:5]:  # 最多5个示例
            if name in self.v1_data and name in self.v2c_data and name in self.v2s_data:
                examples.append({
                    'name': name,
                    'v1': self.v1_data[name],
                    'v2c': self.v2c_data[name],
                    'v2s': self.v2s_data[name]
                })
        
        return examples
    
    def generate_statistics(self) -> Dict:
        """生成统计信息"""
        stats = {
            'v1': {
                'total': len(self.v1_data),
                'by_split': defaultdict(int),
                'amc_count': 0,
                'avg_informal_len': 0,
                'avg_formal_len': 0
            },
            'v2c': {
                'total': len(self.v2c_data),
                'by_split': defaultdict(int),
                'amc_count': 0,
                'avg_informal_len': 0,
                'avg_formal_len': 0
            },
            'v2s': {
                'total': len(self.v2s_data),
                'by_split': defaultdict(int),
                'amc_count': 0,
                'avg_informal_len': 0,
                'avg_formal_len': 0
            }
        }
        
        for version, data in [('v1', self.v1_data), ('v2c', self.v2c_data), ('v2s', self.v2s_data)]:
            total_informal = 0
            total_formal = 0
            
            for name, item in data.items():
                stats[version]['by_split'][item.get('split', 'unknown')] += 1
                
                if 'amc' in name.lower():
                    stats[version]['amc_count'] += 1
                
                informal = item.get('informal_statement', '')
                formal = item.get('formal_statement', '')
                total_informal += len(informal)
                total_formal += len(formal)
            
            stats[version]['avg_informal_len'] = total_informal / len(data) if data else 0
            stats[version]['avg_formal_len'] = total_formal / len(data) if data else 0
        
        return stats
    
    def generate_word_report(self, output_path: str):
        """生成 Word 文档报告"""
        if not DOCX_AVAILABLE:
            print("[ERROR] python-docx 未安装，无法生成 Word 文档")
            return False
        
        doc = Document()
        
        # 标题
        title = doc.add_heading('miniF2F v1 vs v2 数据集详细比较分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加日期
        from datetime import datetime
        date_para = doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 空行
        
        # 1. 执行摘要
        doc.add_heading('1. 执行摘要', 1)
        doc.add_paragraph(
            '本报告详细比较了原始 miniF2F (v1) 数据集与其改进版本 miniF2F_v2 '
            '(包括 v2c 竞赛版本和 v2s 简化版本) 之间的差异。分析包括数据集统计、'
            '结构差异以及具体的示例对比。'
        )
        
        # 2. 数据集统计
        doc.add_heading('2. 数据集统计信息', 1)
        stats = self.generate_statistics()
        
        # 创建统计表格
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '指标'
        header_cells[1].text = 'v1 (原始)'
        header_cells[2].text = 'v2c (竞赛)'
        header_cells[3].text = 'v2s (简化)'
        
        # 数据行
        rows_data = [
            ('总问题数', stats['v1']['total'], stats['v2c']['total'], stats['v2s']['total']),
            ('AMC 问题数', stats['v1']['amc_count'], stats['v2c']['amc_count'], stats['v2s']['amc_count']),
            ('平均非正式陈述长度', f"{stats['v1']['avg_informal_len']:.0f}", 
             f"{stats['v2c']['avg_informal_len']:.0f}", f"{stats['v2s']['avg_informal_len']:.0f}"),
            ('平均正式陈述长度', f"{stats['v1']['avg_formal_len']:.0f}", 
             f"{stats['v2c']['avg_formal_len']:.0f}", f"{stats['v2s']['avg_formal_len']:.0f}"),
        ]
        
        for row_data in rows_data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        # 3. 主要差异
        doc.add_heading('3. 主要差异概述', 1)
        
        differences = [
            ('匹配度', 
             'v1 版本可能存在正式和非正式陈述不匹配的问题',
             'v2 版本确保所有正式和非正式陈述完全匹配'),
            ('选择题处理 - v2c',
             'v1: 直接给出答案',
             'v2c: 包含所有选项，需要证明结果属于选项集合'),
            ('选择题处理 - v2s',
             'v1: 直接给出答案',
             'v2s: 直接给出答案，但陈述更清晰'),
            ('陈述长度',
             f"v1: 平均 {stats['v1']['avg_informal_len']:.0f} 字符",
             f"v2c: {stats['v2c']['avg_informal_len']:.0f} 字符, v2s: {stats['v2s']['avg_informal_len']:.0f} 字符")
        ]
        
        for diff_type, v1_desc, v2_desc in differences:
            doc.add_heading(f'3.{differences.index((diff_type, v1_desc, v2_desc)) + 1} {diff_type}', 2)
            doc.add_paragraph(f'v1: {v1_desc}', style='List Bullet')
            doc.add_paragraph(f'v2: {v2_desc}', style='List Bullet')
        
        # 4. 具体示例对比
        doc.add_heading('4. 具体示例对比', 1)
        examples = self.find_comparison_examples()
        
        for idx, example in enumerate(examples, 1):
            doc.add_heading(f'4.{idx} 示例: {example["name"]}', 2)
            
            # 问题名称和来源
            doc.add_paragraph(f'问题名称: {example["name"]}')
            doc.add_paragraph(f'数据集分割: {example["v1"].get("split", "unknown")}')
            
            # v1 版本
            doc.add_heading('v1 (原始版本)', 3)
            doc.add_paragraph('非正式陈述:')
            v1_informal_text = example['v1'].get('informal_statement', '').strip()
            # 移除开头的 /-- 和结尾的 -/
            if v1_informal_text.startswith('/--'):
                v1_informal_text = v1_informal_text[3:].strip()
            if v1_informal_text.endswith('-/'):
                v1_informal_text = v1_informal_text[:-2].strip()
            v1_informal = doc.add_paragraph(v1_informal_text[:800])
            v1_informal.style = 'Intense Quote'
            
            doc.add_paragraph('正式陈述:')
            v1_formal = doc.add_paragraph(example['v1'].get('formal_statement', '').strip())
            v1_formal.style = 'Intense Quote'
            for run in v1_formal.runs:
                run.font.name = 'Courier New'
            
            # v2c 版本
            doc.add_heading('v2c (竞赛版本)', 3)
            doc.add_paragraph('非正式陈述:')
            v2c_informal_text = example['v2c'].get('informal_statement', '').strip()
            v2c_informal = doc.add_paragraph(v2c_informal_text[:800])
            v2c_informal.style = 'Intense Quote'
            
            doc.add_paragraph('正式陈述:')
            v2c_formal = doc.add_paragraph(example['v2c'].get('formal_statement', '').strip())
            v2c_formal.style = 'Intense Quote'
            for run in v2c_formal.runs:
                run.font.name = 'Courier New'
            
            # v2s 版本
            doc.add_heading('v2s (简化版本)', 3)
            doc.add_paragraph('非正式陈述:')
            v2s_informal_text = example['v2s'].get('informal_statement', '').strip()
            v2s_informal = doc.add_paragraph(v2s_informal_text[:800])
            v2s_informal.style = 'Intense Quote'
            
            doc.add_paragraph('正式陈述:')
            v2s_formal = doc.add_paragraph(example['v2s'].get('formal_statement', '').strip())
            v2s_formal.style = 'Intense Quote'
            for run in v2s_formal.runs:
                run.font.name = 'Courier New'
            
            # 差异分析
            doc.add_heading('差异分析', 3)
            differences_text = self.analyze_example_differences(example)
            doc.add_paragraph(differences_text)
            
            doc.add_page_break()
        
        # 5. 结论
        doc.add_heading('5. 结论', 1)
        doc.add_paragraph(
            '通过对比分析，我们发现 v2 版本相比 v1 版本的主要改进包括：'
        )
        conclusions = [
            '所有正式和非正式陈述完全匹配，提高了数据集的准确性',
            'v2c 版本保持竞赛原样，更适合评估模型在真实竞赛场景下的表现',
            'v2s 版本简化了陈述，更适合训练和基本证明能力评估',
            'v2 版本更准确地反映了数学证明的真实难度'
        ]
        for conclusion in conclusions:
            doc.add_paragraph(conclusion, style='List Bullet')
        
        # 保存文档
        doc.save(output_path)
        print(f"[OK] Word 文档已保存到: {output_path}")
        return True
    
    def analyze_example_differences(self, example: Dict) -> str:
        """分析示例的差异"""
        v1_informal = example['v1'].get('informal_statement', '')
        v2c_informal = example['v2c'].get('informal_statement', '')
        v2s_informal = example['v2s'].get('informal_statement', '')
        
        v1_formal = example['v1'].get('formal_statement', '')
        v2c_formal = example['v2c'].get('formal_statement', '')
        v2s_formal = example['v2s'].get('formal_statement', '')
        
        analysis = []
        
        # 检查是否是选择题
        if 'Show that it is' in v1_informal or 'Prove that it is' in v2c_informal:
            analysis.append("这是一个选择题。")
            analysis.append(f"v1 版本: 直接给出答案，非正式陈述包含 'Show that it is'。")
            
            if 'Prove that it is equal to one of the following options' in v2c_informal:
                analysis.append("v2c 版本: 包含所有选项，需要证明结果属于选项集合。")
            else:
                analysis.append("v2c 版本: 保持竞赛原样，包含所有选项。")
            
            if 'Show that it is' in v2s_informal:
                analysis.append("v2s 版本: 直接给出答案，简化了陈述。")
        
        # 检查正式陈述的差异
        if v1_formal != v2c_formal or v1_formal != v2s_formal:
            analysis.append("\n正式陈述存在差异：")
            if '∨' in v2c_formal and '∨' not in v1_formal:
                analysis.append("- v2c 使用析取 (∨) 表示选项集合")
            if v2s_formal != v2c_formal:
                analysis.append("- v2s 直接证明具体答案，不使用选项集合")
        
        return '\n'.join(analysis) if analysis else "三个版本的陈述基本相同。"
    
    def generate_markdown_report(self, output_path: str):
        """生成 Markdown 报告（备用）"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# miniF2F v1 vs v2 数据集详细比较分析报告\n\n")
            f.write(f"生成时间: {__import__('datetime').datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n\n")
            
            # 统计信息
            stats = self.generate_statistics()
            f.write("## 1. 数据集统计信息\n\n")
            f.write("| 指标 | v1 (原始) | v2c (竞赛) | v2s (简化) |\n")
            f.write("|------|-----------|-------------|-------------|\n")
            f.write(f"| 总问题数 | {stats['v1']['total']} | {stats['v2c']['total']} | {stats['v2s']['total']} |\n")
            f.write(f"| AMC 问题数 | {stats['v1']['amc_count']} | {stats['v2c']['amc_count']} | {stats['v2s']['amc_count']} |\n")
            f.write(f"| 平均非正式陈述长度 | {stats['v1']['avg_informal_len']:.0f} | {stats['v2c']['avg_informal_len']:.0f} | {stats['v2s']['avg_informal_len']:.0f} |\n")
            f.write(f"| 平均正式陈述长度 | {stats['v1']['avg_formal_len']:.0f} | {stats['v2c']['avg_formal_len']:.0f} | {stats['v2s']['avg_formal_len']:.0f} |\n\n")
            
            # 示例
            examples = self.find_comparison_examples()
            f.write("## 2. 具体示例对比\n\n")
            
            for idx, example in enumerate(examples, 1):
                f.write(f"### 2.{idx} 示例: {example['name']}\n\n")
                f.write(f"**问题名称**: {example['name']}\n\n")
                
                f.write("#### v1 (原始版本)\n\n")
                f.write("**非正式陈述**:\n```\n")
                f.write(example['v1'].get('informal_statement', '')[:500])
                f.write("\n```\n\n")
                f.write("**正式陈述**:\n```lean\n")
                f.write(example['v1'].get('formal_statement', ''))
                f.write("\n```\n\n")
                
                f.write("#### v2c (竞赛版本)\n\n")
                f.write("**非正式陈述**:\n```\n")
                f.write(example['v2c'].get('informal_statement', '')[:500])
                f.write("\n```\n\n")
                f.write("**正式陈述**:\n```lean\n")
                f.write(example['v2c'].get('formal_statement', ''))
                f.write("\n```\n\n")
                
                f.write("#### v2s (简化版本)\n\n")
                f.write("**非正式陈述**:\n```\n")
                f.write(example['v2s'].get('informal_statement', '')[:500])
                f.write("\n```\n\n")
                f.write("**正式陈述**:\n```lean\n")
                f.write(example['v2s'].get('formal_statement', ''))
                f.write("\n```\n\n")
                
                f.write("#### 差异分析\n\n")
                f.write(self.analyze_example_differences(example))
                f.write("\n\n---\n\n")
        
        print(f"[OK] Markdown 报告已保存到: {output_path}")


def main():
    """主函数"""
    print("=" * 60)
    print("生成详细比较报告")
    print("=" * 60)
    print()
    
    v1_path = "minif2f_v1/miniF2F_v1.json"
    v2c_path = "../miniF2F_v2/datasets/miniF2F_v2c.json"
    v2s_path = "../miniF2F_v2/datasets/miniF2F_v2s.json"
    
    generator = DetailedReportGenerator(v1_path, v2c_path, v2s_path)
    generator.load_datasets()
    
    # 生成 Word 文档
    if DOCX_AVAILABLE:
        word_path = "miniF2F_comparison_report.docx"
        try:
            if generator.generate_word_report(word_path):
                print(f"[OK] Word 文档已生成: {word_path}")
        except PermissionError:
            print(f"[ERROR] 无法保存文件，可能文件正在被其他程序打开: {word_path}")
            print("请关闭 Word 或其他打开该文件的程序后重试")
        except Exception as e:
            print(f"[ERROR] 生成 Word 文档时出错: {e}")
    else:
        print("[WARN] 无法生成 Word 文档，生成 Markdown 格式")
        generator.generate_markdown_report("detailed_comparison_report.md")
    
    print()
    print("=" * 60)
    print("报告生成完成！")
    print("=" * 60)


if __name__ == "__main__":
    main()
